#!/usr/bin/env python3
"""
Video to Text Converter
Extracts audio from video files, transcribes speech to text, and saves to DOCX format.
"""

import os
import sys
import argparse
from pathlib import Path
from datetime import datetime
from moviepy.editor import VideoFileClip
import speech_recognition as sr
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile


class VideoToTextConverter:
    """Convert video files to text documents."""

    def __init__(self, video_path, output_path=None, language='en-US'):
        """
        Initialize the converter.

        Args:
            video_path (str): Path to the video file
            output_path (str): Path for output DOCX file (optional)
            language (str): Language code for speech recognition (default: en-US)
        """
        self.video_path = Path(video_path)
        if not self.video_path.exists():
            raise FileNotFoundError(f"Video file not found: {video_path}")

        if output_path:
            self.output_path = Path(output_path)
        else:
            # Default output path: same name as video but .docx extension
            self.output_path = self.video_path.with_suffix('.docx')

        self.language = language
        self.recognizer = sr.Recognizer()
        self.temp_audio_path = None

    def extract_audio(self):
        """Extract audio from video file and save as WAV."""
        print(f"Extracting audio from: {self.video_path.name}")

        try:
            # Load video file
            video = VideoFileClip(str(self.video_path))

            # Create temporary WAV file
            temp_fd, self.temp_audio_path = tempfile.mkstemp(suffix='.wav')
            os.close(temp_fd)

            # Extract and save audio
            audio = video.audio
            if audio is None:
                raise ValueError("No audio track found in video file")

            audio.write_audiofile(self.temp_audio_path,
                                 codec='pcm_s16le',
                                 ffmpeg_params=["-ac", "1"],  # mono
                                 logger=None)

            # Close video file
            video.close()
            audio.close()

            print(f"Audio extracted successfully: {self.temp_audio_path}")
            return self.temp_audio_path

        except Exception as e:
            print(f"Error extracting audio: {str(e)}")
            raise

    def transcribe_audio(self, audio_path):
        """
        Transcribe audio file to text using Google Speech Recognition.

        Args:
            audio_path (str): Path to audio file

        Returns:
            str: Transcribed text
        """
        print("Transcribing audio to text...")
        print("This may take a while depending on the video length...")

        full_text = []

        try:
            with sr.AudioFile(audio_path) as source:
                # Get audio file duration
                audio_duration = source.DURATION
                print(f"Audio duration: {audio_duration:.2f} seconds")

                # Process audio in chunks to handle longer files
                chunk_duration = 30  # seconds
                offset = 0
                chunk_num = 1

                while offset < audio_duration:
                    try:
                        # Read chunk
                        audio_chunk = self.recognizer.record(
                            source,
                            duration=min(chunk_duration, audio_duration - offset),
                            offset=offset
                        )

                        print(f"Processing chunk {chunk_num} "
                              f"({offset:.1f}s - {min(offset + chunk_duration, audio_duration):.1f}s)...")

                        # Recognize speech
                        try:
                            text = self.recognizer.recognize_google(
                                audio_chunk,
                                language=self.language,
                                show_all=False
                            )

                            if text:
                                full_text.append(text)
                                print(f"  Transcribed: {text[:50]}...")

                        except sr.UnknownValueError:
                            print(f"  Chunk {chunk_num}: Could not understand audio")
                        except sr.RequestError as e:
                            print(f"  Chunk {chunk_num}: API error - {e}")

                        offset += chunk_duration
                        chunk_num += 1

                    except Exception as e:
                        print(f"Error processing chunk {chunk_num}: {str(e)}")
                        offset += chunk_duration
                        chunk_num += 1
                        continue

        except Exception as e:
            print(f"Error during transcription: {str(e)}")
            raise

        # Combine all text chunks
        transcribed_text = ' '.join(full_text)

        if not transcribed_text:
            print("Warning: No text could be transcribed from the audio")
            transcribed_text = "[No speech detected in audio]"

        return transcribed_text

    def create_docx(self, transcribed_text):
        """
        Create a DOCX document with the transcribed text.

        Args:
            transcribed_text (str): The transcribed text to write
        """
        print(f"Creating DOCX file: {self.output_path}")

        # Create new document
        doc = Document()

        # Add title
        title = doc.add_heading('Video Transcription', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph()
        metadata = doc.add_paragraph()
        metadata.add_run('Source Video: ').bold = True
        metadata.add_run(str(self.video_path.name))

        metadata = doc.add_paragraph()
        metadata.add_run('Transcription Date: ').bold = True
        metadata.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

        metadata = doc.add_paragraph()
        metadata.add_run('Language: ').bold = True
        metadata.add_run(self.language)

        # Add separator
        doc.add_paragraph('_' * 80)

        # Add transcription heading
        doc.add_heading('Transcription Text', level=1)

        # Add transcribed text
        # Split into paragraphs for better readability
        paragraphs = transcribed_text.split('. ')

        for i, para in enumerate(paragraphs):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                if i < len(paragraphs) - 1 and not para.endswith('.'):
                    p.add_run('.')

                # Set font
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

        # Add footer
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        footer = doc.add_paragraph()
        footer.add_run('Generated by Video to Text Converter').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save document
        doc.save(str(self.output_path))
        print(f"DOCX file created successfully: {self.output_path}")

    def convert(self):
        """Main conversion process."""
        try:
            print("\n" + "=" * 60)
            print("VIDEO TO TEXT CONVERTER")
            print("=" * 60 + "\n")

            # Step 1: Extract audio
            audio_path = self.extract_audio()

            # Step 2: Transcribe audio
            transcribed_text = self.transcribe_audio(audio_path)

            # Step 3: Create DOCX
            self.create_docx(transcribed_text)

            print("\n" + "=" * 60)
            print("CONVERSION COMPLETED SUCCESSFULLY!")
            print("=" * 60)
            print(f"\nOutput file: {self.output_path.absolute()}")
            print(f"Transcribed text length: {len(transcribed_text)} characters")

            return self.output_path

        except Exception as e:
            print(f"\nError during conversion: {str(e)}")
            raise

        finally:
            # Clean up temporary audio file
            if self.temp_audio_path and os.path.exists(self.temp_audio_path):
                try:
                    os.remove(self.temp_audio_path)
                    print(f"\nCleaned up temporary file: {self.temp_audio_path}")
                except Exception as e:
                    print(f"\nWarning: Could not delete temporary file: {e}")


def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description='Convert video files to text documents (DOCX format)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python video_to_text.py video.mp4
  python video_to_text.py video.mp4 -o output.docx
  python video_to_text.py video.mp4 -l es-ES
  python video_to_text.py meeting.avi -o meeting_transcript.docx -l en-US

Supported video formats: MP4, AVI, MOV, MKV, WMV, FLV, and more
        """
    )

    parser.add_argument(
        'video',
        help='Path to the video file'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output DOCX file path (default: same name as video with .docx extension)',
        default=None
    )

    parser.add_argument(
        '-l', '--language',
        help='Language code for speech recognition (default: en-US)',
        default='en-US'
    )

    args = parser.parse_args()

    try:
        converter = VideoToTextConverter(
            video_path=args.video,
            output_path=args.output,
            language=args.language
        )
        converter.convert()
        return 0

    except Exception as e:
        print(f"\nFatal error: {str(e)}", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
